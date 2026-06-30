"""
Word 文档生成工具 — docx_tool

适用场景：对外交付的正式研究报告、需要公司模板样式的文件

设计规范：
  - 与项目其他工具（search / rag / sql / excel）保持一致的 _ok / _err 返回结构
  - 延迟 import python-docx，缺少依赖时返回友好错误（参考 search.py 对 langchain-tavily 的处理）
  - @tool(description=...) 提供 LLM 路由决策所需的触发/禁止条件
  - sections / metadata 接受 Any 类型，在函数体内通过 _doc_common 手动校验和容错，
    避免 @tool 装饰器的 Pydantic 校验在函数体外抛出 ValidationError
  - 输入标准化 / 容错逻辑统一收敛到 tools/_doc_common.py

依赖：pip install python-docx
"""

from __future__ import annotations

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any

from langchain_core.tools import tool
from react_agent.utils.tool_helpers import _ok, _err
from react_agent.tools._doc_common import (
    normalize_sections,
    coerce_metadata,
    build_filename,
)

# ── 输出目录（与 excel_tool / md_tool 保持同一约定）────────────────────────
OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "./outputs")).resolve()
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ── 样式常量 ─────────────────────────────────────────────────────────────
_BRAND_BLUE = (0x1F, 0x49, 0x7D)  # 标题 / 表头底色
_HEADER_FILL = "1F497D"  # 表头底色（十六进制字符串，供 OOXML shd 使用）
_ROW_FILL_EVEN = "F2F2F2"  # 数据行偶数行底色
_ROW_FILL_ODD = "FFFFFF"  # 数据行奇数行底色


# ── 辅助函数：安全写入单元格文本并取得可设样式的 run ─────────────────────
def _set_cell(cell, text: str, *, font_size, bold=False, color=None):
    """
    安全地向单元格写入文本并返回首个 run 以便继续设置样式。

    为什么不直接用 `cell.text = text` 后取 `runs[0]`：
      当 text 为空字符串时，python-docx 不会创建任何 run，
      `cell.paragraphs[0].runs[0]` 会抛 IndexError。
      这里显式 add_run（空串也会创建 run），从根上规避该问题。
    """
    from docx.shared import Pt, RGBColor

    para = cell.paragraphs[0]
    # 清空可能存在的默认空 run，保证样式干净
    para.clear()
    run = para.add_run(str(text))
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    return run


def _shade_cell(cell, fill_hex: str):
    """为单元格设置底色（OOXML w:shd 元素）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


# ── 辅助函数：使用 python-docx 构建 Document 对象 ─────────────────────────
def _build_docx(
        title: str,
        sections: list[dict],
        metadata: dict | None = None,
):
    """
    使用 python-docx 构建 Document 对象。调用方负责捕获 ImportError。

    sections 已经过 _doc_common.normalize_sections 净化：
      - level 已是 1~6 的 int（这里再 min(level, 4) 适配 Word 标题层级）
      - bullets 要么是 list 要么是 None
      - table 要么是含 headers/rows 的 dict 要么是 None
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── 页面边距（A4，上 2.5 / 下 2.5 / 左 3.0 / 右 2.5 cm）──────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── 默认正文字体 ──────────────────────────────────────────────────
    doc.styles["Normal"].font.name = "等线"
    doc.styles["Normal"].font.size = Pt(11)

    # ── 文档标题 ────────────────────────────────────────────────────
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_para.runs:  # add_heading 对空标题可能不产生 run，做防御
        title_run = title_para.runs[0]
        title_run.font.color.rgb = RGBColor(*_BRAND_BLUE)
        title_run.font.size = Pt(18)

    # ── 元数据行 ────────────────────────────────────────────────────
    if metadata:
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for k, v in metadata.items():
            run = meta_para.add_run(f"{k}：{v}    ")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()  # 标题后空行

    # ── 章节内容 ────────────────────────────────────────────────────
    for sec in sections:
        # _sanitize_section 钳制到 1~6；Word 内置标题样式仅到 Heading 4，此处再收窄
        level = min(sec.get("level", 2), 4)
        heading = sec.get("heading", "")
        content = sec.get("content", "")
        bullets = sec.get("bullets")
        table = sec.get("table")

        if heading:
            h = doc.add_heading(heading, level=level)
            for run in h.runs:
                run.font.color.rgb = RGBColor(*_BRAND_BLUE)

        # content 内的 \n 拆为独立段落，避免多段文本被挤进单个 Word 段落
        if content:
            for para_text in content.split("\n"):
                doc.add_paragraph(para_text)

        # 列表（bullets 已保证为 list 或 None）
        if bullets:
            for item in bullets:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(item)).font.size = Pt(11)

        # 表格（table 已保证为含 headers/rows 的 dict 或 None）
        if table:
            _render_table(doc, table)

    # ── 页脚 ────────────────────────────────────────────────────────
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("由 ReAct Agent 自动生成    ").font.size = Pt(8)

    return doc


def _render_table(doc, table: dict):
    """渲染单个表格。headers 非法时静默跳过，行列数不匹配时按表头列数对齐。"""
    headers = table.get("headers")
    rows = table.get("rows")

    if not headers or not isinstance(headers, list):
        return
    n_cols = len(headers)

    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = "Table Grid"

    # 表头行：深蓝底白字加粗
    hdr_cells = tbl.rows[0].cells
    for i, h_text in enumerate(headers):
        _set_cell(hdr_cells[i], h_text if h_text is not None else "",
                  font_size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr_cells[i], _HEADER_FILL)

    # 数据行：交替底色，列数按 n_cols 钳制，避免越界 / 缺列
    if isinstance(rows, list):
        for row_idx, row_data in enumerate(rows):
            if not isinstance(row_data, (list, tuple)):
                continue  # 跳过非序列行（如误传的 dict）
            row_cells = tbl.add_row().cells
            fill = _ROW_FILL_EVEN if row_idx % 2 == 0 else _ROW_FILL_ODD
            for col_idx in range(n_cols):
                # 缺列补空串，多余列丢弃，None 渲染为空而非 "None"
                value = row_data[col_idx] if col_idx < len(row_data) else ""
                if value is None:
                    value = ""
                _set_cell(row_cells[col_idx], value, font_size=10)
                _shade_cell(row_cells[col_idx], fill)

    doc.add_paragraph()  # 表格后空行


# ═══════════════════════════════════════════════════════════════════════════
# 工具：docx_tool — Word 文档生成
# ═══════════════════════════════════════════════════════════════════════════
@tool(
    description=(
            "【触发条件】以下情况调用本工具：\n"
            "  - 用户要求生成正式的研究报告、行业分析报告、对外交付文档\n"
            "  - 需要标准排版、公司品牌样式（标题颜色、表格美化、页脚等）\n"
            "  - 内容含有需要精确对齐的多列数据表格\n"
            "  - 最终产物要打印或通过邮件以附件形式发送\n"
            "  - 用户明确要求 .docx / Word 格式\n\n"
            "【不触发条件】以下情况禁止调用本工具：\n"
            "  - 内容仅为轻量备忘或需要在 Markdown 渲染器中使用（应使用 md_tool）\n"
            "  - 用户要求生成 Excel 表格（应使用 excel_tool）\n"
            "  - 内容仅需口头回答，不需要保存为文件\n\n"
            "【输入】\n"
            "  title    : 文档标题，必填，字符串\n"
            "  sections : 章节列表（list[dict]），每个 dict 可含以下字段：\n"
            "             heading(str) / level(int,默认2,范围1-4) / content(str) / "
            "bullets(list[str]) / table(dict)\n"
            "  filename : 文件名（可选，不含扩展名）\n"
            "  mode     : 'timestamp'（默认，不覆盖）或 'overwrite'\n"
            "  metadata : 可选元数据 dict，如 {'作者': '研究团队'}\n\n"
            "【sections 完整示例】\n"
            '  [\n'
            '    {"heading": "一、市场概况", "level": 2, '
            '"content": "2026 年行业整体增长稳健。"},\n'
            '    {"heading": "二、核心数据", "level": 2,\n'
            '     "bullets": ["营收同比 +18%", "毛利率 32%"],\n'
            '     "table": {"headers": ["公司", "毛利率"], '
            '"rows": [["比亚迪", "28%"], ["宁德时代", "25%"]]}}\n'
            '  ]\n\n'
            "【输出】\n"
            "  data.path    — 生成文件的绝对路径\n"
            "  data.size_kb — 文件大小（KB）"
    )
)
def docx_tool(
        title: str = "",
        sections: Any = None,
        filename: str = "",
        mode: str = "timestamp",
        metadata: Any = None,
) -> dict[str, Any]:
    """
    生成 Word（.docx）文档并保存到本地。

    参数：
        title     : 文档标题
        sections  : 章节列表，每个元素为 dict，支持字段：
                    - heading  (str)       章节标题
                    - level    (int, 默认2) 标题级别 1-4
                    - content  (str)       正文段落
                    - bullets  (list[str]) 无序列表项
                    - table    (dict)      {"headers": [...], "rows": [[...], ...]}
        filename  : 文件名（不含扩展名）。为空时使用 title 生成
        mode      : "timestamp"（默认）附加时间戳避免覆盖 / "overwrite" 直接覆盖
        metadata  : 可选元数据 dict，写在标题下方

    返回：
        统一的 _ok / _err 结构
    """
    tool_name = "docx_tool"

    # ════════════════════ 阶段 1：输入校验与标准化 ════════════════════
    t = (str(title) if title else "").strip()
    if not t:
        return _err(
            tool_name=tool_name,
            query=t,
            code="BAD_INPUT",
            message="title 不能为空，请提供文档标题。",
        )

    try:
        normalized = normalize_sections(sections)
    except (ValueError, TypeError, json.JSONDecodeError) as e:
        return _err(
            tool_name=tool_name,
            query=t,
            code="BAD_INPUT",
            message=f"sections 参数格式无法解析：{e}。请传入章节字典列表。",
        )

    if mode not in ("timestamp", "overwrite"):
        mode = "timestamp"  # 容错：非法 mode 默认回退，不拒绝

    metadata = coerce_metadata(metadata)

    # ════════════════════ 阶段 2：依赖检查 ════════════════════
    try:
        import docx  # noqa: F401
    except ImportError:
        return _err(
            tool_name=tool_name,
            query=t,
            code="MISSING_DEPENDENCY",
            message="未安装依赖 python-docx。请执行：pip install python-docx",
        )

    # ════════════════════ 阶段 3：构建并写入 ════════════════════
    try:
        name = build_filename(t, filename, mode, ext="docx")
        out_path = OUTPUT_DIR / name

        meta = {"生成时间": datetime.now().strftime("%Y-%m-%d %H:%M")}
        if metadata:
            meta.update(metadata)

        doc = _build_docx(t, normalized, meta)
        doc.save(str(out_path))

        size_kb = round(out_path.stat().st_size / 1024, 1)
        return _ok(
            tool_name=tool_name,
            query=t,
            data={
                "path": str(out_path.resolve()),
                "size_kb": size_kb,
                "message": f"Word 文档已生成：{name}（{size_kb} KB，{len(normalized)} 个章节）",
            },
        )

    except Exception as e:
        return _err(
            tool_name=tool_name,
            query=t,
            code="WRITE_FAILED",
            message=f"docx_tool 生成失败：{type(e).__name__}: {e}",
        )
